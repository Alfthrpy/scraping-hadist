from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from docx import Document
import re

driver = webdriver.Chrome()
driver.maximize_window()
driver.implicitly_wait(20)

doc = Document()



def cekShigatulAda(arab: str, jml : int):
    result = []
    shigatul = [
    "حَدَّثَنَا", "حَدَّثَنِي", "حَدَّثَتْنَا", "حَدَّثَهُ", "أَخْبَرَنَا", 
    "أَخْبَرَهُ", "فَخَطَبَنَا", "سَمِعْتُ", "أَنْبَأَنَا", "ذَكَرَلَنَا", 
    "قَالَ لَنَا", "عَنْ", "قَرَأْتُ عَلَيْهِ", "قُرِئَ عَلَى", "حَدَّثَنَاأَوْخَبَرَنَا قِرَاءَةًعَلَيْهِ", 
    "الاجازة", "أَجَزْتُلَكَ", "أَجَزْتُ", "هَذَاسَمَاعِى", "اَنْبَأَنَا", 
    "نَاوَلَنَا", "أَعْلَمَنِى", "اَوْصَىاِلَىَّ", "قَرَاْتُ", "أَنَّ", "سَمِعَ", "إِنَّ", 
    "قَالَ قَالَ ", "أَخْبَرَهُ", "رَفَعَهُ","حَدَّثَتْنِي", "ذَكَرْنَا","أَخْبَرَنِي","أَخْبَرَتْنِي"
    ]


    arab_words = re.findall(r'\S+', arab)

    for word in arab_words:
        if word in shigatul:
            result.append(word)

        if len(result) == jml:
            break
    
    return result




def ekstrak(awal,akhir,sheet):
    if awal == 0:
        awal = 1
    jml = awal*10
    jml = jml - 10
    for i in range(awal,akhir+1):
        
        driver.get(f"https://hadits.tazkia.ac.id/hadits/buku/1?page_haditses={i}")

        hadist_hadist = WebDriverWait(driver, 5).until(
                    EC.presence_of_all_elements_located((By.CLASS_NAME, "hadits"))
                )
        # Tunggu hingga elemen pencarian dapat diakses
        
        for iterate in range(0,len(hadist_hadist)):
            list_rawi = []
            list_thobaqoh = []
            i = hadist_hadist[iterate]
            jml_rawi = 0
            
            head = i.find_element(By.CLASS_NAME,"mb-3")
            doc.add_paragraph("")  # Baris kosong
            doc.add_paragraph(head.text)
            doc.add_paragraph("")  # Baris kosong
            

            btn = i.find_element(By.TAG_NAME,"button")
            driver.execute_script("arguments[0].click();", btn)

            modal = driver.find_element(By.ID, "haditsModal")
            a = modal.find_elements(By.TAG_NAME, "a")
            print(len(a))
            href_list = []

            

            for i in range(0, len(a)):
                current_a = a[i]
                try:
                    rawi = WebDriverWait(current_a, 5).until(
                        EC.presence_of_all_elements_located((By.TAG_NAME, "div"))
                    )
                    rawi = current_a.find_element(By.TAG_NAME, "div")
                    time.sleep(1)
                    list_rawi.append(rawi.text)
                    jml_rawi = len(list_rawi)
                    print(rawi.text)


                    # Simpan nilai href sebagai identitas elemen
                    href_value = current_a.get_attribute("href")
                    href_list.append(href_value)


                except:
                    break


            for href_value in href_list:
                driver.get(href_value)

                thobqoh = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "tr"))
                )
                list_thobaqoh.append(thobqoh.text)
                print(thobqoh.text)


                driver.back()

            hadist_hadist = WebDriverWait(driver, 5).until(
                EC.presence_of_all_elements_located((By.CLASS_NAME, "hadits"))
            )
            
            i = hadist_hadist[iterate]
            arab = i.find_element(By.CLASS_NAME, "arabic")
            doc.add_paragraph(arab.text)

            terjemah = i.find_element(By.CLASS_NAME,"indonesia")
            doc.add_paragraph(terjemah.text)


            shigatul = cekShigatulAda(arab.text,jml_rawi)
            print(shigatul)

            print()

            jml += 1
            
            list_rawi.reverse()
            list_thobaqoh.reverse()
            print(list_rawi)
            print(list_thobaqoh)
            doc.add_paragraph("")  # Baris kosong
            table = doc.add_table(rows=1, cols=5)
            table.style = 'TableGrid'  # Gunakan gaya tabel dengan grid

            columns = table.rows[0].cells
            columns[0].text = 'No'
            columns[1].text = 'Rawi'
            columns[2].text = 'Shigatul Ada'
            columns[3].text = 'Thobaqoh'
            columns[4].text = 'Kualitas'

            for i in range(0, len(list_rawi)):
                try:
                    row_cells = table.add_row().cells
                    if i == 0:
                        row_cells[0].text = str(jml)
                    row_cells[1].text = list_rawi[i]
                    row_cells[2].text = shigatul[i]
                    row_cells[3].text = list_thobaqoh[i]
                    if i == 0:
                        row_cells[4].text = "Shahih"
                except:
                    row_cells = table.add_row().cells
                    if i == 0:
                        row_cells[0].text = str(jml)
                    row_cells[1].text = list_rawi[i]
                    row_cells[2].text = "SHIGAT TIDAK DITEMUKAN, MOHON UNTUK DICARI SECARA MANUAL!"
                    row_cells[3].text = list_thobaqoh[i]
                    if i == 0:
                        row_cells[4].text = "Shahih"
            

    doc.save('safta.docx')



        



awal = 1110
akhir = 1150

ekstrak(int(awal/10),int(akhir/10),doc)
    # Tutup browser setelah selesai
driver.quit()
